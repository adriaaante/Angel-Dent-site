#!/usr/bin/env python3
"""Единый договор клиники с пациентом → .docx (для правки и печати) + .pdf.

Текст живёт в `contract_text.py` — там же объяснено, почему договор один на
три клиники и что обязательно заполнить перед печатью.

    python3 _materials/dogovor-pacienta/build.py

.docx нужен клинике и юристу — его правят и печатают. .pdf собирается тем же
текстом, чтобы владелец мог посмотреть документ с телефона: там docx теряет
таблицы и отступы (та же причина, что и у счетов в _materials/docs/).
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (BaseDocTemplate, Frame, KeepTogether, PageTemplate,
                                Paragraph, Spacer)

HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(HERE))
import contract_text as T  # noqa: E402

OUT = HERE / "out"
NAME = "Договор-на-оказание-стоматологических-услуг"
FONTS = Path("/tmp/claude-0/-home-user/0c942a14-b9d6-5314-9918-7c11e119dac0/scratchpad/fonts")

INK = colors.HexColor("#111418")
MUTED = colors.HexColor("#5D6D77")


def placeholders(text: str) -> str:
    """{{...}} — то, что заполняют руками. В обоих форматах показываем одинаково."""
    return text.replace("{{", "___ ").replace("}}", " ___")


# ────────────────────────────── DOCX ──────────────────────────────
# Вёрстка «как у юриста»: поля под подшивку, висячий отступ у номеров пунктов,
# заголовки не отрываются от текста, таблицы для мест оказания, согласий и
# подписей, в колонтитуле — номер страницы и подписи сторон (защита от
# подмены листа).
BODY_FONT = "Times New Roman"
BODY_SIZE = 10.5


def _set_font(run, size=None, bold=False, italic=False):
    run.font.name = BODY_FONT
    run.font.size = Pt(size or BODY_SIZE)
    run.bold = bold
    run.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)


def docx_base() -> Document:
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = BODY_FONT
    st.font.size = Pt(BODY_SIZE)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = st.paragraph_format
    pf.space_after = Pt(3)
    pf.line_spacing = 1.1
    for s in doc.sections:
        s.top_margin = Cm(1.8)
        s.bottom_margin = Cm(1.6)
        s.left_margin = Cm(2.5)      # поле под подшивку
        s.right_margin = Cm(1.5)
    return doc


def p(doc, text="", bold=False, align="just", size=None, space=None, indent=None,
      first_line=None, keep_with_next=False, italic=False):
    par = doc.add_paragraph()
    par.alignment = {"just": WD_ALIGN_PARAGRAPH.JUSTIFY, "left": WD_ALIGN_PARAGRAPH.LEFT,
                     "center": WD_ALIGN_PARAGRAPH.CENTER,
                     "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    pf = par.paragraph_format
    if space is not None:
        pf.space_after = Pt(space)
    if indent is not None:
        pf.left_indent = Cm(indent)
    if first_line is not None:
        pf.first_line_indent = Cm(first_line)
    pf.keep_with_next = keep_with_next
    _set_font(par.add_run(placeholders(text)), size, bold, italic)
    return par


def clause(doc, num: str, text: str):
    """Пункт договора: номер висит слева, текст выровнен по общей левой кромке."""
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = par.paragraph_format
    pf.left_indent = Cm(1.1)
    pf.first_line_indent = Cm(-1.1)   # висячий отступ
    pf.space_after = Pt(3)
    _set_font(par.add_run(f"{num}\t"), bold=True)
    _set_font(par.add_run(placeholders(text)))
    par.paragraph_format.tab_stops.add_tab_stop(Cm(1.1))
    return par


def borderless(table):
    tbl = table._tbl
    borders = tbl.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tbl.makeelement(qn(f"w:{edge}"), {qn("w:val"): "none", qn("w:sz"): "0"})
        borders.append(el)
    tbl.tblPr.append(borders)
    return table


def set_widths(table, widths_cm):
    """python-docx ширину берёт с ячеек — ставим и на них, и на колонки."""
    table.autofit = False
    for row in table.rows:
        row.cant_split = True
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)
    for col, w in zip(table.columns, widths_cm):
        col.width = Cm(w)
    return table


def cell_text(cell, lines, bold_first=False, size=None):
    cell.text = ""
    for i, line in enumerate(lines):
        par = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        par.paragraph_format.space_after = Pt(1)
        _set_font(par.add_run(placeholders(line)), size, bold=bold_first and i == 0)


def add_footer(doc):
    """Номер страницы и строка подписей — чтобы лист нельзя было подменить."""
    from docx.oxml import OxmlElement
    footer = doc.sections[0].footer
    par = footer.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(2)
    _set_font(par.add_run("Исполнитель ____________   Заказчик ____________   "
                          "Пациент ____________          Страница "), 8)
    for tag, txt in (("begin", None), (None, "PAGE"), ("end", None)):
        if tag:
            el = OxmlElement("w:fldChar"); el.set(qn("w:fldCharType"), tag)
        else:
            el = OxmlElement("w:instrText"); el.set(qn("xml:space"), "preserve")
            el.text = " PAGE "
        run = par.add_run(); _set_font(run, 8); run._r.append(el)
    _set_font(par.add_run(" из "), 8)
    for tag, txt in (("begin", None), (None, "NUMPAGES"), ("end", None)):
        if tag:
            el = OxmlElement("w:fldChar"); el.set(qn("w:fldCharType"), tag)
        else:
            el = OxmlElement("w:instrText"); el.set(qn("xml:space"), "preserve")
            el.text = " NUMPAGES "
        run = par.add_run(); _set_font(run, 8); run._r.append(el)


def build_docx() -> Path:
    doc = docx_base()
    add_footer(doc)

    for line in T.TITLE.split("\n"):
        p(doc, line, bold=True, align="center", size=12.5, space=2, keep_with_next=True)
    p(doc, "", space=4)

    # Город слева, дата справа — одной строкой через таблицу без границ.
    head = borderless(doc.add_table(rows=1, cols=2))
    head.autofit = True
    cell_text(head.rows[0].cells[0], ["г. ____________________"])
    par = head.rows[0].cells[1].paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_font(par.add_run("«____» ______________ 20___ г."))
    p(doc, "", space=6)

    for block in T.PREAMBLE:
        p(doc, block, first_line=0.75)

    p(doc, "", space=2)
    p(doc, T.PLACE_BLOCK_TITLE, bold=True, align="left", space=3, keep_with_next=True)
    place = set_widths(borderless(doc.add_table(rows=len(T.CLINICS), cols=2)),
                       [1.4, 15.6])
    for row, (name, addr, phone) in zip(place.rows, T.CLINICS):
        cell_text(row.cells[0], ["[     ]"])
        cell_text(row.cells[1], [f"Стоматология {name} — {addr}, тел. {phone}"])
    p(doc, "", space=2)
    p(doc, T.PLACE_NOTE, size=9, space=8, italic=True)

    for title, items in T.SECTIONS:
        p(doc, title, bold=True, align="left", space=3, size=11, keep_with_next=True)
        for num, text in items:
            clause(doc, num, text)
        p(doc, "", space=2)

    p(doc, "14. Приложения", bold=True, align="left", size=11, space=3,
      keep_with_next=True)
    for a in T.APPENDICES:
        p(doc, a, align="left", space=1, indent=1.1, first_line=-0.5)
    p(doc, "", space=6)

    p(doc, T.CONSENTS_TITLE, bold=True, align="left", space=3, size=11,
      keep_with_next=True)
    cons = set_widths(borderless(doc.add_table(rows=len(T.CONSENTS), cols=2)),
                      [4.4, 12.6])
    for row, c in zip(cons.rows, T.CONSENTS):
        cell_text(row.cells[0], ["[   ] Согласен(на)", "[   ] Не согласен(на)"])
        cell_text(row.cells[1], [c])
    p(doc, "", space=8)

    # ── Реквизиты и подписи: отдельным листом, чтобы блок не рвался ────────
    doc.add_page_break()
    p(doc, "15. Реквизиты и подписи Сторон", bold=True, align="left", size=11,
      space=4, keep_with_next=True)
    c = T.COMPANY
    sign = set_widths(borderless(doc.add_table(rows=1, cols=2)), [8.5, 8.5])
    cell_text(sign.rows[0].cells[0], [
        "ИСПОЛНИТЕЛЬ",
        c["full"],
        f"Место нахождения: {c['address']}",
        f"ОГРН {c['ogrn']}   ИНН {c['inn']}   КПП {c['kpp']}",
        "Расчётный счёт ______________________",
        "Банк ________________________________",
        "Корр. счёт __________________________",
        "БИК ____________________",
        f"Телефон {c['phone']}",
        "Электронная почта ____________________",
        f"Лицензия № {c['license']} (бессрочно)",
        "",
        "Должность ___________________________",
        "",
        "____________ / ____________________  М. П.",
        "         (подпись)                (расшифровка)",
    ], bold_first=True)
    cell_text(sign.rows[0].cells[1], [
        "ЗАКАЗЧИК",
        "Ф. И. О. ____________________________",
        "____________________________________",
        "Дата рождения ______________________",
        "Паспорт: серия ______ № _____________",
        "выдан _______________________________",
        "____________________________________",
        "дата выдачи ________________________",
        "Адрес места жительства ______________",
        "____________________________________",
        "Иной адрес для ответов ______________",
        "Телефон ____________________",
        "Электронная почта ____________________",
        "",
        "______________ / ______________________",
        "         (подпись)                (расшифровка)",
    ], bold_first=True)
    p(doc, "", space=8)

    p(doc, "ПАЦИЕНТ — потребитель услуг (заполняется, если Пациент и Заказчик — "
           "разные лица)", bold=True, align="left", space=3, keep_with_next=True)
    pat = set_widths(borderless(doc.add_table(rows=1, cols=2)), [8.5, 8.5])
    cell_text(pat.rows[0].cells[0], [
        "Ф. И. О. ____________________________",
        "____________________________________",
        "Дата рождения ______________________",
        "Паспорт: серия ______ № _____________",
        "выдан _______________________________",
        "дата выдачи ________________________",
    ])
    cell_text(pat.rows[0].cells[1], [
        "Адрес места жительства ______________",
        "____________________________________",
        "Телефон ____________________________",
        "Заказчик приходится Пациенту ________",
        "(родитель, опекун, попечитель, иное)",
        "Документ о полномочиях представителя",
        "____________________________________",
    ])
    p(doc, "", space=6)
    p(doc, "______________ / ______________________", align="left", space=1)
    p(doc, "         (подпись)                (расшифровка)", align="left", space=3,
      size=9)
    p(doc, "За Пациента, не достигшего 15 лет, Договор подписывает законный "
           "представитель (ст. 54 Федерального закона от 21.11.2011 № 323-ФЗ). "
           "Экземпляр Договора получен каждой Стороной.", size=9, italic=True)

    path = OUT / f"{NAME}.docx"
    doc.save(path)
    return path


# ────────────────────────────── PDF ──────────────────────────────
def pdf_styles() -> dict:
    pdfmetrics.registerFont(TTFont("Onest", str(FONTS / "Onest-Regular.ttf")))
    pdfmetrics.registerFont(TTFont("Onest-Bold", str(FONTS / "Onest-Bold.ttf")))

    def st(name, **kw):
        base = dict(fontName="Onest", fontSize=8.8, leading=12.2, textColor=INK,
                    alignment=TA_JUSTIFY)
        base.update(kw)
        return ParagraphStyle(name, **base)

    return {
        "title": st("title", fontName="Onest-Bold", fontSize=13, leading=17,
                    alignment=TA_CENTER, spaceAfter=2),
        "meta": st("meta", alignment=TA_LEFT, spaceAfter=8),
        "h": st("h", fontName="Onest-Bold", fontSize=10, leading=13.5,
                alignment=TA_LEFT, spaceBefore=8, spaceAfter=3),
        "p": st("p", spaceAfter=3.5),
        "li": st("li", alignment=TA_LEFT, spaceAfter=2, leftIndent=10),
        "small": st("small", fontSize=7.8, leading=10.5, textColor=MUTED, spaceAfter=4),
        "sign": st("sign", alignment=TA_LEFT, spaceAfter=2),
    }


def build_pdf() -> Path:
    s = pdf_styles()
    path = OUT / f"{NAME}.pdf"
    doc = BaseDocTemplate(str(path), pagesize=A4,
                          leftMargin=20 * mm, rightMargin=14 * mm,
                          topMargin=15 * mm, bottomMargin=15 * mm,
                          title="Договор на оказание стоматологических услуг",
                          author=T.COMPANY["short"])
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="f")

    def decorate(canvas, d):
        canvas.saveState()
        canvas.setFont("Onest", 7)
        canvas.setFillColor(MUTED)
        canvas.drawRightString(A4[0] - doc.rightMargin, 9 * mm, f"стр. {d.page}")
        canvas.restoreState()

    doc.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=decorate)])

    P = lambda t, k="p": Paragraph(placeholders(t), s[k])
    story = []
    for line in T.TITLE.split("\n"):
        story.append(P(line, "title"))
    story.append(Spacer(1, 4))
    story.append(P("г. ____________________&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;"
                   "«____» ______________ 20___ г.", "meta"))
    for block in T.PREAMBLE:
        story.append(P(block))

    story.append(P(T.PLACE_BLOCK_TITLE, "h"))
    for name, addr, phone in T.CLINICS:
        story.append(P(f"[&nbsp;&nbsp;&nbsp;]&nbsp; Стоматология {name} — {addr}, тел. {phone}", "li"))
    story.append(P(T.PLACE_NOTE, "small"))

    for title, items in T.SECTIONS:
        block = [P(title, "h")] + [P(f"{n} {t}") for n, t in items]
        story.append(KeepTogether(block) if len(items) <= 4 else block[0])
        if len(items) > 4:
            story.extend(block[1:])

    story.append(P("14. Приложения", "h"))
    for a in T.APPENDICES:
        story.append(P(a, "li"))

    story.append(P(T.CONSENTS_TITLE, "h"))
    for c in T.CONSENTS:
        story.append(P(f"[&nbsp;&nbsp;&nbsp;] Согласен(на)&nbsp;&nbsp;&nbsp; [&nbsp;&nbsp;&nbsp;] Не согласен(на)&nbsp;&nbsp;—&nbsp;&nbsp;{c}", "li"))

    story.append(P("15. Реквизиты и подписи Сторон", "h"))
    c = T.COMPANY
    for line in ("<b>ИСПОЛНИТЕЛЬ</b>", c["full"],
                 f"Место нахождения и фактический адрес: {c['address']}",
                 f"ОГРН {c['ogrn']} · ИНН {c['inn']} · КПП {c['kpp']}",
                 "Расчётный счёт ____________________ в ____________________",
                 "Корреспондентский счёт ____________________ · БИК ______________",
                 f"Телефон {c['phone']} · Электронная почта ____________________",
                 f"Лицензия № {c['license']} (бессрочно)", "",
                 "_______________ / ________________________&nbsp;&nbsp; М. П.", "",
                 "<b>ЗАКАЗЧИК</b>",
                 "Ф. И. О. ______________________________________________",
                 "Дата рождения ____________________",
                 "Паспорт: серия ______ № ______________ выдан ____________________",
                 "Адрес места жительства ______________________________________",
                 "Иной адрес для ответов на обращения __________________________",
                 "Телефон ____________________ · Электронная почта ______________", "",
                 "_______________ / ________________________", "",
                 "<b>ПАЦИЕНТ</b> (потребитель услуг; заполняется, если Пациент и "
                 "Заказчик — разные лица)",
                 "Ф. И. О. ______________________________________________",
                 "Дата рождения ____________________",
                 "Паспорт: серия ______ № ______________ выдан ____________________",
                 "Адрес места жительства ______________________________________",
                 "Телефон ____________________",
                 "Заказчик приходится Пациенту ____________________ "
                 "(родитель, опекун, попечитель, иное)",
                 "Документ о полномочиях законного представителя (свидетельство о "
                 "рождении, акт органа опеки, доверенность) ____________________", "",
                 "_______________ / ________________________"):
        story.append(P(line or "&nbsp;", "sign"))
    story.append(P("За Пациента, не достигшего 15 лет, Договор подписывает законный "
                   "представитель (ст. 54 Федерального закона № 323-ФЗ).", "small"))

    doc.build(story)
    return path


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    d = build_docx()
    f = build_pdf()
    for x in (d, f):
        print(f"  ✓ {x.relative_to(HERE.parent.parent)}  ({x.stat().st_size // 1024} КБ)")
    print("\n  ⚠️ Перед печатью обязательно заполнить:")
    for t in T.TODO:
        print(f"     — {t}")


if __name__ == "__main__":
    main()
